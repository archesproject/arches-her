# coding: utf-8
'''
ARCHES - a program developed to inventory and manage immovable cultural heritage.
Copyright (C) 2013 J. Paul Getty Trust and World Monuments Fund

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU Affero General Public License as
published by the Free Software Foundation, either version 3 of the
License, or (at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
GNU Affero General Public License for more details.

You should have received a copy of the GNU Affero General Public License
along with this program. If not, see <http://www.gnu.org/licenses/>.
'''

import json
import os
import uuid
from datetime import datetime
import docx
import textwrap
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from html.parser import HTMLParser
from html.entities import name2codepoint
from django.core.files.uploadedfile import UploadedFile
from django.http import HttpRequest, HttpResponseNotFound
from django.utils.translation import ugettext as _
from django.views.generic import View
from arches.app.datatypes.datatypes import DataTypeFactory
from arches.app.models import models
from arches.app.models.concept import ConceptValue
from arches.app.models.resource import Resource
from arches.app.models.system_settings import settings
from arches.app.models.tile import Tile
from arches.app.utils.response import JSONResponse
from arches.app.views.tile import TileData


class FileTemplateView(View):

    def __init__(self):
        self.doc = None
        self.resource = None


    def get(self, request):
        parenttile_id = request.GET.get('parenttile_id')
        parent_tile = Tile.objects.get(tileid=parenttile_id)
        letter_tiles = Tile.objects.filter(parenttile=parent_tile)
        file_list_node_id = "96f8830a-8490-11ea-9aba-f875a44e0e11" # Digital Object
        url = None
        for tile in letter_tiles:
            if url is not None:
                break
            for data_obj in tile.data[file_list_node_id]:
                if data_obj['status'] == 'uploaded':
                    url = data_obj['url']
                    break

        if url is not None:
            return JSONResponse({'msg':'success','download':url })
        return HttpResponseNotFound("No letters tile matching query by parent tile")
    
    
    def post(self, request): 
        datatype_factory = DataTypeFactory()
        template_id = request.POST.get('template_id')
        parenttile_id = request.POST.get('parenttile_id')
        resourceinstance_id = request.POST.get('resourceinstance_id', None)
        transaction_id = request.POST.get('transaction_id', uuid.uuid1())
        self.resource = Resource.objects.get(resourceinstanceid=resourceinstance_id)
        self.resource.load_tiles()

        template_name = self.get_template_path(template_id)
        template_path = os.path.join(settings.APP_ROOT, 'docx', template_name)

        if os.path.exists(os.path.join(settings.APP_ROOT, 'uploadedfiles','docx')) is False:
            os.mkdir(os.path.join(settings.APP_ROOT, 'uploadedfiles','docx'))

        try:
            self.doc = Document(template_path)
        except:
            return HttpResponseNotFound("No Template Found")

        self.edit_letter(self.resource, datatype_factory)

        date = datetime.today()
        date = date.strftime("%Y")+'-'+date.strftime("%m")+'-'+date.strftime("%d")
        new_file_name = date+'_'+template_name
        new_file_path = os.path.join(settings.APP_ROOT, 'uploadedfiles/docx', new_file_name)

        new_req = HttpRequest()
        new_req.method = 'POST'
        new_req.user = request.user
        new_req.POST['data'] = None
        host = request.get_host()

        self.doc.save(new_file_path)
        saved_file = open(new_file_path, 'rb')
        stat = os.stat(new_file_path)
        file_data = UploadedFile(saved_file)
        file_list_node_id = "96f8830a-8490-11ea-9aba-f875a44e0e11" # Digital Object
 
        tile = json.dumps({
            "tileid":None,
            "data": {
                file_list_node_id: [{
                    "name":new_file_name,
                    "accepted":True,
                    "height":0,
                    "lastModified":stat.st_mtime,
                    "size":stat.st_size,
                    "status":"queued",
                    "type":"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "width":0,
                    "url":None,
                    "file_id":None,
                    "index":0,
                    "content":"blob:"+host+"/{0}".format(uuid.uuid4())
                }]
            },
            "nodegroup_id":"7db68c6c-8490-11ea-a543-f875a44e0e11",
            "parenttile_id":None,
            "resourceinstance_id":"",
            "sortorder":0,
            "tiles":{}
        })

        new_req = HttpRequest()
        new_req.method = 'POST'
        new_req.user = request.user
        new_req.POST['data'] = tile
        new_req.POST['transaction_id'] = transaction_id
        new_req.FILES['file-list_' + file_list_node_id] = file_data
        new_tile = TileData()
        new_tile.action = "update_tile"

        response = TileData.post(new_tile, new_req)
        if response.status_code == 200:
            tile = json.loads(response.content)
            return JSONResponse({'tile':tile, 'status':'success' })

        return HttpResponseNotFound(response.status_code)


    def get_template_path(self, template_id):
        template_dict = { # keys are valueids from "Letters" concept list; values are known file names
            "01dec356-e72e-40e6-b1b1-b847b9799d2f": "No progression letter.docx", # Letter A
            "320abc26-db82-44a6-be11-8d44aaa23365":	"No Need to Consult letter.docx", # Letter A2
            "fd15c6c7-e94d-4914-8d51-a98bda6f4a7b":	"Pre-app Predetermination letter.docx", # Letter B1
            "8cc91474-11ce-47d9-b886-f0e3fc49d277": "Predetermination Letter.docx", # Letter B2
            "08bb630d-a27b-45bc-a13f-567b428018c5": "Condition Two-Stage Letter.docx", # Letter C
            "e14bd058-e9f2-48f8-8ef5-337310c3420f":	"Pre-App Recommend Condition Letter.docx", # Letter D1
            "92e745c3-7157-4831-bce9-73792d32abec":	"Condition Investigation Letter.docx", # Letter D2
            "41f3d0bb-a94d-469f-80c8-85ab03283972":	"Condition Historic Building Recording Letter.docx", # Letter D3
            "7f1e7061-8bb0-4338-9342-118f1e9214aa":	"WSI Approval Letter.docx", # Letter F1
            "eaa8a075-50e6-4c3d-ac08-fbe84865f577":	"WSI Amend Letter.docx", # Letter F2
            "8d605e5c-d0da-4b72-9ce3-2f7dac3381d1":	"Post Excavation Assessment Approval Letter.docx", # Letter G - PXA Approval
            "a31061ea-9b80-435f-82c8-94dc10afcbae":	"Condition Satisfied Letter.docx", # Letter H
            "eed24dd2-85a0-4402-a6ba-bda426b5da89":	"Blank Adviser Letter.docx", # Letter I - Bespoke Letter
            # No template available yet
            "a26c77ff-1d04-4b76-a45f-417f7ed24333": "", # Additional Condition Text
            "8c12a812-8000-4ec9-913d-c6fd516117f2": "", # Archaeological Recommendation Text
            # No concept selection available
            "missing 0": "Conditions Scope Notes.docx",
            "missing 1": "Mitigation Scope Notes.docx",
        }
        for key, value in list(template_dict.items()):
            if key == template_id:
                return value

        return None

    def edit_letter(self, consultation, datatype_factory):
        template_dict = {
            "Reference": "c128bf36-9384-11ea-bfe1-f875a44e0e11",
            "Primary Reference Number": "b37552bd-9527-11ea-97f4-f875a44e0e11",
            "Casework Officer": "4ea4a197-184f-11eb-9152-f875a44e0e11",
            "Completion Date": "40eff4ce-893a-11ea-ae2e-f875a44e0e11",
            "Consultation Name": "4ad69684-951f-11ea-b5c3-f875a44e0e11",
            "Proposal Description": "1b0e15ec-8864-11ea-8493-f875a44e0e11",
            "Log Date": "40eff4cd-893a-11ea-b0cc-f875a44e0e11",
            "Signature": "4ea4a197-184f-11eb-9152-f875a44e0e11",
            "Archaeological Priority Area": "58a2b98f-a255-11e9-9a30-00224800b26d",
            "Assessment of Significance": "8d41e4e0-a250-11e9-877f-00224800b26d",
            "Condition Type": "56fa335d-06fa-11eb-8328-f875a44e0e11",
            "Condition": "c36808b0-952c-11ea-9ff0-f875a44e0e11",
            "Mitigation Type": "e2585f8a-51a3-11eb-a7be-f875a44e0e11",
            "Mitigation": "bfd39106-51a3-11eb-9104-f875a44e0e11"
        }

        self.replace_in_letter(consultation.tiles, template_dict, datatype_factory)    

    def replace_in_letter(self, tiles, template_dict, datatype_factory):
        mapping_dict = {
            "Reference": "",
            "Primary Reference Number": "",
            "Casework Officer": "",
            "Completion Date": "",
            "Consultation Name": "",
            "Proposal Description": "",
            "Log Date": "",
            "Signature": "",
            "Archaeological Priority Area": "",
            "Assessment of Significance": "",
            "Condition": "",
            "Mitigation": "",
            "Casework Officer Email": "",
            "Casework Officer Number": "",
            "Contact Name": "",
            "Address of consulting organisation": "",
            "Name of person consulting": "",
        }

        def get_value_from_tile(tile, node_id):
            current_node = models.Node.objects.get(nodeid=node_id)
            datatype = datatype_factory.get_instance(current_node.datatype)
            returnvalue = datatype.get_display_value(tile, current_node)
            return "" if returnvalue is None else returnvalue

        # Advice and Conditions.
        advice_nodegroup_id = '8d41e49f-a250-11e9-b6b3-00224800b26d'
        advice_node_id = 'c36808b0-952c-11ea-9ff0-f875a44e0e11'
        advice_type_node_id = '56fa335d-06fa-11eb-8328-f875a44e0e11'
        conditions = []

        # Action and Mitigations.
        action_nodegroup_id = 'a5e15f5c-51a3-11eb-b240-f875a44e0e11'
        action_node_id = 'bfd39106-51a3-11eb-9104-f875a44e0e11'
        action_type_node_id = 'e2585f8a-51a3-11eb-a7be-f875a44e0e11'
        mitigations = []

        mitigation_scope_dict = {}
        mitigation_notes_path = os.path.join(settings.APP_ROOT, "docx/Mitigation Scope Notes.json")
        with open(mitigation_notes_path, "rb") as openfile:
            mitigation_scope_dict = json.loads(openfile.read())

        for tile in tiles:
            mitigation = {}
            condition = {}
            if str(tile.nodegroup_id) == action_nodegroup_id:
                mitigation_scopenote = mitigation_scope_dict.get(mitigation_scope_dict.get(get_value_from_tile(tile, action_type_node_id), ""), "")
                if len(mitigation_scopenote) > 0:
                    mitigation_scopenote = "<br>" + mitigation_scopenote
                mitigation["content"] = "<p>{}</p><p>{}</p>".format(get_value_from_tile(tile, action_node_id), mitigation_scopenote)
                mitigation["type"] = get_value_from_tile(tile, action_type_node_id)
            elif str(tile.nodegroup_id) == advice_nodegroup_id:
                condition["content"] = "<p>{}</p>".format(get_value_from_tile(tile, advice_node_id))
                template_name = self.get_template_path(self.request._post['template_id'])
                if template_name == "WSI Amend Letter.docx" or template_name == "WSI Approval Letter.docx":
                    condition["type"] = ""
                else:
                    condition["type"] = get_value_from_tile(tile, advice_type_node_id)
            else:
                for key, value in list(template_dict.items()):
                    if value in tile.data:
                        lookup_val = get_value_from_tile(tile, value)
                        try:
                            mapping_dict[key] = lookup_val
                        except TypeError:
                            pass
            if len(mitigation) > 0:
                mitigations.append(mitigation)
            elif len(condition) > 0:
                conditions.append(condition)

            contactNodeId = "b7304f4c-3ace-11eb-8884-f875a44e0e11"
            contacts = {
                "Applicant": "4ea4a19a-184f-11eb-aef8-f875a44e0e11",
                "Planning Officer": "4ea4a192-184f-11eb-a0d6-f875a44e0e11",
                "Owner": "4ea4c885-184f-11eb-b4d5-f875a44e0e11",
                "Agent": "4ea4c884-184f-11eb-b64d-f875a44e0e11",
                "Casework Officer": "4ea4a197-184f-11eb-9152-f875a44e0e11"
            }

            addressNodegroupId = "5f93048e-80a9-11ea-b0da-f875a44e0e11"
            nameNodegroupId = "4110f741-1a44-11e9-885e-000d3ab1e588"
            contactDetailsNodegroupId = "2547c12f-9505-11ea-a507-f875a44e0e11"
            contactNameForCorrespondenceNodeId = "2beefb56-4084-11eb-bcc5-f875a44e0e11"
            fullnameNodeId = "5f8ded26-7ef9-11ea-8e29-f875a44e0e11"
            firstnameNodeId = "2caeb5e7-7b44-11ea-a919-f875a44e0e11"
            lastnameNodeId = "96a3942a-7e53-11ea-8b5a-f875a44e0e11"
            nameTitleNodeId = "6da2f03b-7e55-11ea-8fe5-f875a44e0e11"
            nameUseTypeNodeId = "4110f747-1a44-11e9-96b7-000d3ab1e588"
            forCorrespondenceNameValueId = "85c26c81-e356-4454-a2ba-67e7ad9b95cd"
            primaryNameValueId = "2df285fa-9cf2-45e7-bc05-a67b7d7ddc2f"
            contactPointNodeId = "2547c133-9505-11ea-8e49-f875a44e0e11"
            contactPointTypeNodeId = "2547c132-9505-11ea-b22f-f875a44e0e11"
            contactPointTypeMailValueId = "e6d433a2-7f77-4eb7-96f2-57ebe0ac251e"
            addressDict = {
                "Building Name": "b3a2761d-effb-11eb-9867-a87eeabdefba",
                "Building Number": "b3a2761f-effb-11eb-9059-a87eeabdefba",
                "Street": "b3a27621-effb-11eb-83e6-a87eeabdefba",
                "Locality": "b3a28c1a-effb-11eb-a811-a87eeabdefba",
                "Town or City": "b3a27617-effb-11eb-a80f-a87eeabdefba",
                "Postcode": "b3a27619-effb-11eb-a66d-a87eeabdefba",
            }
            if contactNodeId in tile.data:
                caseAgentResourceId = None
                contactResourceiId = None
                if tile.data[contacts["Casework Officer"]]:
                    caseAgentResourceId = tile.data[contacts["Casework Officer"]][0]["resourceId"]
                if caseAgentResourceId:
                    caseAgentResource = Resource.objects.get(resourceinstanceid=caseAgentResourceId)
                    caseAgentResource.load_tiles()
                    for caseAgentTile in caseAgentResource.tiles:
                        if caseAgentTile.nodegroup.nodegroupid == uuid.UUID(contactDetailsNodegroupId):
                            if caseAgentTile.data[contactPointTypeNodeId] == "0f466b8b-a347-439f-9b61-bee9811ccbf0":
                                mapping_dict["Casework Officer Email"] = caseAgentTile.data[contactPointNodeId]
                            elif caseAgentTile.data[contactPointTypeNodeId] == "75e6cfad-7418-4ed3-841b-3c083d7df30b":
                                mapping_dict["Casework Officer Number"] = caseAgentTile.data[contactPointNodeId]

                if tile.data[contactNodeId] == "5cc97bfd-d76f-40fc-be60-fbb9dfb28fc4" and tile.data[contacts["Planning Officer"]]:
                    contactResourceiId = tile.data[contacts["Planning Officer"]][0]["resourceId"]
                elif tile.data[contactNodeId] == "d88aa873-848c-45cb-b967-4febe7397912" and tile.data[contacts["Owner"]]:
                    contactResourceiId = tile.data[contacts["Owner"]][0]["resourceId"]
                elif tile.data[contactNodeId] == "dcaf8850-9cfc-44ea-9fd4-0ca419806e2b" and tile.data[contacts["Agent"]]:
                    contactResourceiId = tile.data[contacts["Agent"]][0]["resourceId"]
                
                if contactResourceiId:
                    contactResource = Resource.objects.get(resourceinstanceid=contactResourceiId)
                    contactResource.load_tiles()

                    for contactTile in contactResource.tiles:
                        if contactTile.nodegroup.nodegroupid == uuid.UUID(nameNodegroupId):
                            if mapping_dict["Name of person consulting"] == "" or contactTile.data[nameUseTypeNodeId] == primaryNameValueId:
                                nameTitle = ConceptValue(contactTile.data[nameTitleNodeId]).value
                                fullName = "{0} {1}".format(get_value_from_tile(contactTile, firstnameNodeId), get_value_from_tile(contactTile, lastnameNodeId))
                                mapping_dict["Name of person consulting"] = "{0} {1}".format(nameTitle, fullName) if nameTitle else fullName
                        elif contactTile.nodegroup.nodegroupid == uuid.UUID(contactDetailsNodegroupId):
                            if contactTile.data[contactPointTypeNodeId] == contactPointTypeMailValueId:
                                mapping_dict["Contact Name"] = contactTile.data[contactNameForCorrespondenceNodeId]
                                addressConsult = get_value_from_tile(contactTile, contactPointNodeId).replace(", ", "<br>").replace(",", "<br>")
                                mapping_dict["Address of consulting organisation"] = addressConsult

        for mitigation in mitigations:
            mapping_dict["Mitigation"] += "<b>{}</b>{}<br><br>".format(mitigation["type"], mitigation["content"])
        
        for condition in conditions:
            mapping_dict["Condition"] += "<b>{}</b>{}<br><br>".format(condition["type"], condition["content"])

        associate_heritage = mapping_dict["Archaeological Priority Area"]
        if associate_heritage == "":
            mapping_dict["Archaeological Priority Area"] = "The planning application is not in an area of archaeological interest."
        else:
            mapping_dict["Archaeological Priority Area"] = "The planning application lies in an area of archaeological interest (Archaeological Priority Area) identified in the Local Plan: {}".format(associate_heritage)
            
        for key in mapping_dict:
            html = False
            if '<' in mapping_dict[key]: # look for html tag, not ideal
                html = True
            self.replace_string(self.doc, key, mapping_dict[key], html)

    def replace_string(self, document, key, v, html=False):
        # Note that the intent here is to preserve how things are styled in the docx
        # easiest way is to iterate through p.runs, not as fast as iterating through parent.paragraphs
        # advantage of the former is that replacing run.text preserves styling, replacing p.text does not
        
        def parse_html_to_docx(p, k, v):
            style = p.style
            if k in p.text:
                p.clear()
                document_html_parser = DocumentHTMLParser(p, document)
                document_html_parser.insert_into_paragraph_and_feed(v)

        
        def replace_in_runs(p_list, k, v):
            for paragraph in p_list:
                if html is True:
                    parse_html_to_docx(paragraph, k, v)
                for i, run in enumerate(paragraph.runs):
                    if k in run.text: # now check if html
                        run_style = run.style
                        run.text = run.text.replace(k, v)
                    elif i == (len(paragraph.runs) - 1) and k in paragraph.text: # backstop case: rogue text outside of run obj - must fix template
                        paragraph.text = paragraph.text.replace(k, v)

        def iterate_tables(t_list, k, v):
            for table in t_list:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_runs(cell.paragraphs, k, v)
        
        if v is not None and key is not None:
            k = "<"+key+">"
            doc = document
            # some of these are probably unnecessary
            # foot_style = styles['Footer']
            # head_style = styles['Header']
            # t_style = None
            # p_style = None
            run_style = None

            if len(doc.paragraphs) > 0:
                replace_in_runs(doc.paragraphs, k, v)

            if len(doc.tables) > 0:
                iterate_tables(doc.tables, k, v)
            
            if len(doc.sections) > 0:
                for section in doc.sections:
                    replace_in_runs(section.footer.paragraphs, k, v)
                    iterate_tables(section.footer.tables, k, v)
                    replace_in_runs(section.header.paragraphs, k, v)
                    iterate_tables(section.header.tables, k, v)

    
    def insert_image(self, document, k, v, image_path=None, config=None):
        # going to need to write custom logic depending on how images should be placed/styled

        return True


    def insert_custom(self, document, k, v, config=None):
        # perhaps replaces {{custom_object}} with pre-determined text structure with custom style/format

        return True

        
class DocumentHTMLParser(HTMLParser):
    def __init__(self, paragraph, document):
        HTMLParser.__init__(self)
        self.document = document
        self.paragraph = paragraph
        self.table = None
        self.table_cols = 0
        self.table_rows = 0
        self.max_cols_reached = False
        self.td_cursor = False
        self.hyperlink = False
        self.list_style = "ul"
        self.ol_counter = 1
        self.run = self.paragraph.add_run()

    def insert_paragraph_after(self, paragraph, text=None, style=None):
            """Insert a new paragraph after the given paragraph."""
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            if text:
                new_para.add_run(text)
            if style is not None:
                new_para.style = style
            return new_para

    def add_hyperlink(self, paragraph, url, text, color=None, underline=None):
            # This gets access to the document.xml.rels file and gets a new relation id value
            part = self.paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            # Create the w:hyperlink tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

            # Create a w:r element
            new_run = docx.oxml.shared.OxmlElement('w:r')

            # Create a new w:rPr element
            rPr = docx.oxml.shared.OxmlElement('w:rPr')

            # Add color if it is given
            if not color is None:
                c = docx.oxml.shared.OxmlElement('w:color')
                c.set(docx.oxml.shared.qn('w:val'), color)
                rPr.append(c) # #5384da ; rgb(83,132,218)
            
            # Remove underlining if it is requested
            if not underline:
                u = docx.oxml.shared.OxmlElement('w:u')
                u.set(docx.oxml.shared.qn('w:val'), 'none')
                rPr.append(u)

            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

            return hyperlink

    def insert_into_paragraph_and_feed(self, html):
        self.run = self.paragraph.add_run()
        self.feed(html)

    def handle_starttag(self, tag, attrs):
        self.run = self.paragraph.add_run()
        if tag == "i" or tag == "em":
            self.run.italic = True
        if tag == "b" or tag == "strong":
            self.run.bold = True
        if tag == "s":
            self.run.strike = True
        if tag == "u":
            self.run.underline = True
        if tag == "ol":
            self.list_style = "ol"
        if tag == "ul":
            self.list_style = "ul"
        if tag in ["br", "ul", "ol"]:
            self.run.add_break()
        if tag == "li":
            if self.list_style == 'ul':
                self.run.add_text('● ')
            else:
                self.run.add_text(str(self.ol_counter)+'. ')
                self.ol_counter += 1
        if tag == "p":
            self.run.add_break()
            # self.run.add_break()
            # self.run.add_tab()
        if tag == "a":
            self.hyperlink = attrs[0][1]
        if tag == "table":
            self.table = self.document.add_table(self.table_rows, self.table_cols)
            self.table.autofit = True
        if tag == "tr":
            self.table_rows+= 1
            self.table.add_row()
        if tag == "td":
            self.table_cols+= 1
            if self.max_cols_reached is False:
                self.table.add_column(1)
            self.td_cursor = True

    def handle_endtag(self, tag):
        if tag in ["br", "li", "ul", "ol"]:
            self.run.add_break()
        self.run = self.paragraph.add_run()
        if tag == "ol":
            self.ol_counter = 1
        if tag == "table":
            tbl = self.table._tbl
            p = self.paragraph._p
            p.addnext(tbl)
            self.table = None
            self.table_cols = 0
            self.table_rows = 0
        if tag == "tr":
            self.table_cols = 0
            self.max_cols_reached = True
        if tag == "td":
            self.td_cursor = False

    def handle_data(self, data):
        if "&#39;" in data:
            data = data.replace("&#39;","\'")

        if self.hyperlink is not False:
            blue = docx.shared.RGBColor(83,132,218)
            color = blue.__str__()
            self.add_hyperlink(self.paragraph, self.hyperlink, data, color)
            self.hyperlink = False
        elif self.td_cursor is True:
            self.table.cell(self.table_rows-1, self.table_cols-1).add_paragraph(data) #formatting?
        else:
            self.run.add_text(data)

    def handle_entityref(self, name):
        c = chr(name2codepoint[name])
        self.run.add_text(c)

    def handle_charref(self, name):
        if name.startswith('x'):
            c = chr(int(name[1:], 16))
        else:
            c = chr(int(name))
        self.run.add_text(c)

