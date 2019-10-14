# coding: utf-8
'''
ARCHES - a program developed to inventory and manage immovable cultural heritage.
Copyright (C) 2013 J. Paul Getty Trust and World Monuments Fund

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU Affero General Public License as
published by the Free Software Foundation, either version 3 of the
License, or (at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
GNU Affero General Public License for more details.

You should have received a copy of the GNU Affero General Public License
along with this program. If not, see <http://www.gnu.org/licenses/>.
'''

import json
import os
import uuid
from datetime import datetime
import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from html.parser import HTMLParser
from html.entities import name2codepoint
from pprint import pprint
from django.core.files.uploadedfile import UploadedFile
from django.http import HttpRequest, HttpResponseNotFound
from django.utils.translation import ugettext as _
from django.views.generic import View
from arches.app.datatypes.datatypes import DataTypeFactory
from arches.app.models import models
from arches.app.models.resource import Resource
from arches.app.models.system_settings import settings
from arches.app.models.tile import Tile
from arches.app.utils.betterJSONSerializer import JSONSerializer, JSONDeserializer
from arches.app.utils.response import JSONResponse
from arches.app.views.tile import TileData


class FileTemplateView(View):

    def __init__(self):
        self.doc = None
        self.resource = None


    def get(self, request):
        parenttile_id = request.GET.get('parenttile_id')
        parent_tile = Tile.objects.get(tileid=parenttile_id)
        letter_tiles = Tile.objects.filter(parenttile=parent_tile)
        file_list_node_id = "8d41e4d1-a250-11e9-9a12-00224800b26d"
        url = None
        for tile in letter_tiles:
            if url is not None:
                break
            for data_obj in tile.data[file_list_node_id]:
                if data_obj['status'] == 'uploaded':
                    url = data_obj['url']
                    break

        if url is not None:
            return JSONResponse({'msg':'success','download':url })
        return HttpResponseNotFound("No letters tile matching query by parent tile")
    
    
    def post(self, request): 
        # data = JSONDeserializer().deserialize(request.body)
        datatype_factory = DataTypeFactory()
        template_id = request.POST.get('template_id')
        parenttile_id = request.POST.get('parenttile_id')
        resourceinstance_id = request.POST.get('resourceinstance_id', None)
        self.resource = Resource.objects.get(resourceinstanceid=resourceinstance_id)
        self.resource.load_tiles()

        template_name = self.get_template_path(template_id)
        template_path = os.path.join(settings.APP_ROOT, 'docx', template_name)

        if os.path.exists(os.path.join(settings.APP_ROOT, 'uploadedfiles','docx')) is False:
            os.mkdir(os.path.join(settings.APP_ROOT, 'uploadedfiles','docx'))

        self.doc = Document(template_path)

        if template_name == 'GLAAS Planning Letter A - No Progression - template.docx':
            self.edit_letter_A(self.resource, datatype_factory)
        elif template_name == 'GLAAS Planning Letter B2 - Predetermination - template.docx':
            self.edit_letter_B2(self.resource, datatype_factory)

        date = datetime.today()
        date = date.strftime("%Y")+'-'+date.strftime("%m")+'-'+date.strftime("%d")
        new_file_name = date+'_'+template_name
        new_file_path = os.path.join(settings.APP_ROOT, 'uploadedfiles/docx', new_file_name)

        new_req = HttpRequest()
        new_req.method = 'POST'
        new_req.user = request.user
        new_req.POST['data'] = None
        host = request.get_host()

        self.doc.save(new_file_path)
        saved_file = open(new_file_path, 'rt')
        stat = os.stat(new_file_path)
        file_data = UploadedFile(saved_file)
        file_list_node_id = "8d41e4d1-a250-11e9-9a12-00224800b26d"

        tile = json.dumps({
            "tileid":None,
            "data": {
                file_list_node_id: [{
                    "name":new_file_name,
                    "accepted":True,
                    "height":0,
                    "lastModified":stat.st_mtime,
                    "size":stat.st_size,
                    "status":"queued",
                    "type":"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "width":0,
                    "url":None,
                    "file_id":None,
                    "index":0,
                    "content":"blob:"+host+"/{0}".format(uuid.uuid4())
                }]
            },
            "nodegroup_id":"8d41e4d1-a250-11e9-9a12-00224800b26d",
            "parenttile_id":parenttile_id,
            "resourceinstance_id":resourceinstance_id,
            "sortorder":0,
            "tiles":{}
        })

        new_req = HttpRequest()
        new_req.method = 'POST'
        new_req.user = request.user
        new_req.POST['data'] = tile
        new_req.FILES['file-list_' + file_list_node_id] = file_data
        new_tile_data_instance = TileData()

        post_resp = TileData.post(new_tile_data_instance, new_req)

        if post_resp.status_code == 200:
            return JSONResponse({'tile':tile, 'status':'success' })

        return HttpResponseNotFound(post_resp.status_code)


    def get_template_path(self, template_id):
        template_dict = { # keys are conceptids from "Letters" concept list; values are known file names
            "a26c77ff-1d04-4b76-a45f-417f7ed24333":'', # Addit Cond Text
            "8c12a812-8000-4ec9-913d-c6fd516117f2":'', # Arch Rec Text
            "01dec356-e72e-40e6-b1b1-b847b9799d2f":'GLAAS Planning Letter A - No Progression - template.docx', # Letter A
            "320abc26-db82-44a6-be11-8d44aaa23365":'', # Letter A2
            "fd15c6c7-e94d-4914-8d51-a98bda6f4a7b":'', # Letter B1
            "8cc91474-11ce-47d9-b886-f0e3fc49d277":'GLAAS Planning Letter B2 - Predetermination - template.docx', # Letter B2
            "08bb630d-a27b-45bc-a13f-567b428018c5":'GLAAS Planning Letter C - Condition two stage - template.docx' # Letter C
            }
        for key, value in list(template_dict.items()):
            if key == template_id:
                return value

        return None


    """
    def edit_letter_X(self, consultation, datatype_factory):
        # dict of string/node pairs specific to docx template
        template_dict = {
           string_key1: 'node_id1',
           string_key2: 'node_id2'
        }
        self.replace_in_letter(consultation.tiles, template_dict, datatype_factory)
    """


    def edit_letter_A(self, consultation, datatype_factory):
        template_dict = {
            'Case Officer':'8d41e4d4-a250-11e9-a3ff-00224800b26d',
            'Completion Date': '8d41e4cd-a250-11e9-a25b-00224800b26d',
            'Proposal': '8d41e4bd-a250-11e9-89e8-00224800b26d',
            'Log Date': '8d41e4cf-a250-11e9-a86d-00224800b26d',
            'Action': 'caf5bff8-a3d7-11e9-a37c-00224800b26d'
        }
        self.replace_in_letter(consultation.tiles, template_dict, datatype_factory)

    
    def edit_letter_B2(self, consultation, datatype_factory):
        template_dict = {
            'Case Officer':'8d41e4d4-a250-11e9-a3ff-00224800b26d',
            'Completion Date': '8d41e4cd-a250-11e9-a25b-00224800b26d',
            'Proposal': '8d41e4bd-a250-11e9-89e8-00224800b26d',
            'Log Date': '8d41e4cf-a250-11e9-a86d-00224800b26d',
            'Action': 'caf5bff8-a3d7-11e9-a37c-00224800b26d',
            'Site Name': '???'
        }
        self.replace_in_letter(consultation.tiles, template_dict, datatype_factory)


    def replace_in_letter(self, tiles, template_dict, datatype_factory):
        for tile in tiles:
            for key, value in list(template_dict.items()):
                html = False
                if value in tile.data:
                    my_node = models.Node.objects.get(nodeid=value)
                    datatype = datatype_factory.get_instance(my_node.datatype)
                    lookup_val = datatype.get_display_value(tile, my_node)
                    if '<' in lookup_val: # not ideal
                        html = True
                    self.replace_string(self.doc, key, lookup_val, html)

    
    def replace_string(self, document, key, v, html=False):
        # Note that the intent here is to preserve how things are styled in the docx
        # easiest way is to iterate through p.runs, not as fast as iterating through parent.paragraphs
        # advantage of the former is that replacing run.text preserves styling, replacing p.text does not
        
        def parse_html_to_docx(p, k, v):
            style = p.style
            if k in p.text:
                p.clear()
                document_html_parser = DocumentHTMLParser(p, document)
                document_html_parser.insert_into_paragraph_and_feed(v)

        
        def replace_in_runs(p_list, k, v):
            for paragraph in p_list:
                if html is True:
                    parse_html_to_docx(paragraph, k, v)
                for i, run in enumerate(paragraph.runs):
                    if k in run.text: # now check if html
                        run_style = run.style
                        run.text = run.text.replace(k, v)
                    elif i == (len(paragraph.runs) - 1) and k in paragraph.text: # backstop case: rogue text outside of run obj - must fix template
                        paragraph.text = paragraph.text.replace(k, v)

        def iterate_tables(t_list, k, v):
            for table in t_list:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_runs(cell.paragraphs, k, v)
        
        if v is not None and key is not None:
            k = "{{"+key+"}}"
            doc = document
            # some of these are probably unnecessary
            # foot_style = styles['Footer']
            # head_style = styles['Header']
            # t_style = None
            # p_style = None
            run_style = None

            if len(doc.paragraphs) > 0:
                replace_in_runs(doc.paragraphs, k, v)

            if len(doc.tables) > 0:
                iterate_tables(doc.tables, k, v)
            
            if len(doc.sections) > 0:
                for section in doc.sections:
                    replace_in_runs(section.footer.paragraphs, k, v)
                    iterate_tables(section.footer.tables, k, v)
                    replace_in_runs(section.header.paragraphs, k, v)
                    iterate_tables(section.header.tables, k, v)

    
    def insert_image(self, document, k, v, image_path=None, config=None):
        # going to need to write custom logic depending on how images should be placed/styled

        return True


    def insert_custom(self, document, k, v, config=None):
        # perhaps replaces {{custom_object}} with pre-determined text structure with custom style/format

        return True

        
class DocumentHTMLParser(HTMLParser):
    def __init__(self, paragraph, document):
        HTMLParser.__init__(self)
        self.document = document
        self.paragraph = paragraph
        self.table = None
        self.table_cols = 0
        self.table_rows = 0
        self.max_cols_reached = False
        self.td_cursor = False
        self.hyperlink = False
        self.list_style = "ul"
        self.ol_counter = 1
        self.run = self.paragraph.add_run()

    def insert_paragraph_after(self, paragraph, text=None, style=None):
            """Insert a new paragraph after the given paragraph."""
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            if text:
                new_para.add_run(text)
            if style is not None:
                new_para.style = style
            return new_para

    def add_hyperlink(self, paragraph, url, text, color=None, underline=None):
            # This gets access to the document.xml.rels file and gets a new relation id value
            part = self.paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            # Create the w:hyperlink tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

            # Create a w:r element
            new_run = docx.oxml.shared.OxmlElement('w:r')

            # Create a new w:rPr element
            rPr = docx.oxml.shared.OxmlElement('w:rPr')

            # Add color if it is given
            if not color is None:
                c = docx.oxml.shared.OxmlElement('w:color')
                c.set(docx.oxml.shared.qn('w:val'), color)
                rPr.append(c) # #5384da ; rgb(83,132,218)
            
            # Remove underlining if it is requested
            if not underline:
                u = docx.oxml.shared.OxmlElement('w:u')
                u.set(docx.oxml.shared.qn('w:val'), 'none')
                rPr.append(u)

            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

            return hyperlink

    def insert_into_paragraph_and_feed(self, html):
        self.run = self.paragraph.add_run()
        self.feed(html)

    def handle_starttag(self, tag, attrs):
        # print(tag,attrs)
        self.run = self.paragraph.add_run()
        if tag == "i" or tag == "em":
            self.run.italic = True
        if tag == "b" or tag == "strong":
            self.run.bold = True
        if tag == "s":
            self.run.strike = True
        if tag == "u":
            self.run.underline = True
        if tag == "ol":
            self.list_style = "ol"
        if tag == "ul":
            self.list_style = "ul"
        if tag in ["br", "ul", "ol"]:
            self.run.add_break()
        if tag == "li":
            if self.list_style == 'ul':
                self.run.add_text('● ')
            else:
                self.run.add_text(str(self.ol_counter)+'. ')
                self.ol_counter += 1
        if tag == "p":
            self.run.add_break()
            # self.run.add_break()
            # self.run.add_tab()
        if tag == "a":
            self.hyperlink = attrs[0][1]
        if tag == "table":
            self.table = self.document.add_table(self.table_rows, self.table_cols)
            self.table.autofit = True
        if tag == "tr":
            self.table_rows+= 1
            self.table.add_row()
        if tag == "td":
            self.table_cols+= 1
            if self.max_cols_reached is False:
                self.table.add_column(1)
            self.td_cursor = True

    def handle_endtag(self, tag):
        if tag in ["br", "li", "ul", "ol"]:
            self.run.add_break()
        self.run = self.paragraph.add_run()
        if tag == "ol":
            self.ol_counter = 1
        if tag == "table":
            tbl = self.table._tbl
            p = self.paragraph._p
            p.addnext(tbl)
            self.table = None
            self.table_cols = 0
            self.table_rows = 0
        if tag == "tr":
            self.table_cols = 0
            self.max_cols_reached = True
        if tag == "td":
            self.td_cursor = False

    def handle_data(self, data):
        if "&#39;" in data:
            data = data.replace("&#39;","\'")

        if self.hyperlink is not False:
            blue = docx.shared.RGBColor(83,132,218)
            color = blue.__str__()
            self.add_hyperlink(self.paragraph, self.hyperlink, data, color)
            self.hyperlink = False
        elif self.td_cursor is True:
            self.table.cell(self.table_rows-1, self.table_cols-1).add_paragraph(data) #formatting?
        else:
            self.run.add_text(data)

    def handle_entityref(self, name):
        c = chr(name2codepoint[name])
        self.run.add_text(c)

    def handle_charref(self, name):
        if name.startswith('x'):
            c = chr(int(name[1:], 16))
        else:
            c = chr(int(name))
        self.run.add_text(c)

